import docx
import os

def extrair_texto_docx(caminho_arquivo):
    doc = docx.Document(caminho_arquivo)
    texto_completo = []
    
    for paragrafo in doc.paragraphs:
        texto_completo.append(paragrafo.text)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto_completo.append(paragrafo.text)
    
    return '\n'.join(texto_completo)

if __name__ == "__main__":
    caminho_arquivo = r"C:\Users\usuario\Documents\VICENTE\Atividades Práticas de Testes Automatizados com Selenium\Atividade Prática de Testes Automatizados com Selenium.docx"
    
    if not os.path.exists(caminho_arquivo):
        print(f"Arquivo não encontrado: {caminho_arquivo}")
    else:
        texto_extraido = extrair_texto_docx(caminho_arquivo)
        
        with open("conteudo_atividade.txt", "w", encoding="utf-8") as arquivo_saida:
            arquivo_saida.write(texto_extraido)
        
        print("Conteúdo extraído e salvo em 'conteudo_atividade.txt'")

